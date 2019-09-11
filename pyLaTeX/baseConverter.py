#!/usr/bin/env python3
import os, re

docStart = '\\begin{document}'
docEnd   = '\\end{document}'

sections = { re.compile( r'\\section{(.+)}' ) : 1 }
subs     = ( ('\&', '&'), )

from docx import Document
class baseConverter( object ):
  def __init__(self, TeX_file):
    super().__init__();
    self.TeX_file  = os.path.realpath( TeX_file );
    self.docx_file = '.'.join( self.TeX_file.split('.')[:-1] ) + '.docx'
    self.docx      = None
  ##############################################################################
  def convert(self):
    self.docx   = Document();
    with open( self.TeX_file, 'r' ) as fid:
      data = fid.read();
    
    for sub in subs:
      data = data.replace( sub[0], sub[1] )

    preamble, text  = data.split( docStart )
    text, postamble = text.split( docEnd )
    
    for line in text.splitlines():
      if (len(line) > 0) and (line[0] != '%'):
        if not self.checkSection( line ):
          self.docx.add_paragraph( line )
        
    
    self.docx.save( self.docx_file )

  ##############################################################################
  def checkSection( self, line ):
    for key, val in sections.items():
      section = key.findall( line )
      if len( section ) == 1:
        self.docx.add_heading( section[0], level = val )
        return True;
    return False;
    

if __name__ == "__main__":
  import sys
  inst = baseConverter( sys.argv[1] );
  inst.convert()
  
    